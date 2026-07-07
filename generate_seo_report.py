"""
SEO Audit Report Generator
===========================
Crawls a given URL, performs a comprehensive on-page SEO audit,
and generates a professional Word (.docx) report.

Usage:
    python generate_seo_report.py
"""

import re
import json
import time
import os
from datetime import datetime
from urllib.parse import urljoin, urlparse, urldefrag
from collections import deque

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Configuration ─────────────────────────────────────────────────────────────
TARGET_URL = "https://www.sevenunique.com/"
MAX_PAGES = 15
MAX_DEPTH = 3
TIMEOUT = 20
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

# ── SEO Benchmarks ────────────────────────────────────────────────────────────
BENCHMARKS = {
    "title_length": {"min": 50, "max": 60, "label": "50-60 characters"},
    "meta_desc_length": {"min": 120, "max": 160, "label": "120-160 characters"},
    "h1_count": {"target": 1, "label": "Exactly 1 H1 tag"},
    "word_count": {"min": 300, "label": "300+ words"},
    "internal_links": {"min": 3, "label": "3+ internal links"},
    "image_alt": {"target": 100, "label": "100% alt text coverage"},
}


# ══════════════════════════════════════════════════════════════════════════════
#                            CRAWLER
# ══════════════════════════════════════════════════════════════════════════════

def normalize_url(url):
    normalized, _ = urldefrag(url.strip())
    parsed = urlparse(normalized)
    path = parsed.path or "/"
    if path != "/" and path.endswith("/"):
        path = path.rstrip("/")
    return parsed._replace(netloc=parsed.netloc.lower(), path=path).geturl()


def same_domain(base_domain, url):
    candidate = urlparse(url).netloc.lower()
    base = base_domain.lower()
    return candidate == base or candidate.endswith(f".{base}") or base.endswith(f".{candidate}")


def extract_page_data(response, depth, base_domain):
    """Parse a page and extract all SEO-relevant data."""
    soup = BeautifulSoup(response.text, "lxml")
    final_url = normalize_url(str(response.url))

    # Title
    title = soup.title.get_text(strip=True) if soup.title else ""

    # Meta Description
    meta_tag = soup.find("meta", attrs={"name": lambda v: v and v.lower() == "description"})
    description = (meta_tag.get("content", "").strip()) if meta_tag else ""

    # Meta Robots
    robots_tag = soup.find("meta", attrs={"name": lambda v: v and v.lower() == "robots"})
    robots_content = (robots_tag.get("content", "").strip().lower()) if robots_tag else ""

    # Viewport
    viewport_tag = soup.find("meta", attrs={"name": lambda v: v and v.lower() == "viewport"})
    has_viewport = bool(viewport_tag)

    # Canonical
    canonical_tag = soup.find("link", rel=lambda v: v and "canonical" in str(v).lower())
    canonical_url = ""
    if canonical_tag and canonical_tag.get("href"):
        canonical_url = normalize_url(urljoin(final_url, canonical_tag["href"]))

    # Headings
    headings = {}
    for i in range(1, 7):
        tag = f"h{i}"
        headings[tag] = [h.get_text(" ", strip=True) for h in soup.find_all(tag)]

    # Word count
    exclude_tags = {"script", "style", "noscript", "svg", "header", "footer", "nav", "head", "form"}
    texts = soup.find_all(string=True)
    visible_texts = []
    for t in texts:
        content = str(t).strip()
        if not content:
            continue
        parent = t.parent
        if not parent:
            continue
        is_excluded = False
        curr = parent
        while curr and curr.name != '[document]':
            if curr.name in exclude_tags:
                is_excluded = True
                break
            curr = curr.parent
        if not is_excluded:
            visible_texts.append(content)
    combined = " ".join(visible_texts)
    word_count = len(re.findall(r"\b[\w'-]+\b", combined))

    # Images
    images = soup.find_all("img")
    total_images = len(images)
    missing_alt = [img.get("src", "inline") for img in images if not (img.get("alt") or "").strip()]

    # Links
    internal_links = []
    external_links = []
    external_domains = set()
    dofollow = 0
    nofollow = 0

    for a in soup.find_all("a", href=True):
        href = urljoin(final_url, a["href"])
        parsed = urlparse(href)
        if parsed.scheme not in {"http", "https"}:
            continue
        rel = " ".join(a.get("rel", [])).lower()
        if "nofollow" in rel:
            nofollow += 1
        else:
            dofollow += 1
        if same_domain(base_domain, href):
            internal_links.append(normalize_url(href))
        else:
            external_links.append(href)
            external_domains.add(parsed.netloc)

    # Structured data
    schema_types = []
    for script in soup.find_all("script", attrs={"type": lambda v: v and "ld+json" in v.lower()}):
        raw = script.string or script.get_text(strip=True)
        if not raw:
            continue
        try:
            parsed_json = json.loads(raw)
            stack = parsed_json if isinstance(parsed_json, list) else [parsed_json]
            while stack:
                item = stack.pop()
                if isinstance(item, dict):
                    t = item.get("@type")
                    if isinstance(t, str):
                        schema_types.append(t)
                    elif isinstance(t, list):
                        schema_types.extend(v for v in t if isinstance(v, str))
                    graph = item.get("@graph")
                    if isinstance(graph, list):
                        stack.extend(graph)
        except:
            pass

    # Open Graph
    og_title = ""
    og_desc = ""
    og_image = ""
    og_title_tag = soup.find("meta", property="og:title")
    if og_title_tag:
        og_title = og_title_tag.get("content", "")
    og_desc_tag = soup.find("meta", property="og:description")
    if og_desc_tag:
        og_desc = og_desc_tag.get("content", "")
    og_img_tag = soup.find("meta", property="og:image")
    if og_img_tag:
        og_image = og_img_tag.get("content", "")

    # Twitter Card
    twitter_card = ""
    tc_tag = soup.find("meta", attrs={"name": "twitter:card"})
    if tc_tag:
        twitter_card = tc_tag.get("content", "")

    # Page type classification
    path = urlparse(final_url).path.lower().strip("/")
    if not path:
        page_type = "Homepage"
    elif any(p in path for p in ["contact", "get-in-touch"]):
        page_type = "Contact/Conversion"
    elif any(p in path for p in ["about", "company", "team"]):
        page_type = "About/Trust"
    elif any(p in path for p in ["blog", "news", "article"]):
        page_type = "Blog/Resource"
    elif any(p in path for p in ["service", "solution", "product"]):
        page_type = "Service/Product"
    else:
        page_type = "Other"

    is_indexable = "noindex" not in robots_content

    return {
        "url": final_url,
        "title": title,
        "description": description,
        "robots_content": robots_content,
        "has_viewport": has_viewport,
        "canonical_url": canonical_url,
        "headings": headings,
        "word_count": word_count,
        "total_images": total_images,
        "missing_alt_images": missing_alt,
        "internal_links": internal_links,
        "external_links": external_links,
        "external_domains": list(external_domains),
        "dofollow_links": dofollow,
        "nofollow_links": nofollow,
        "schema_types": schema_types,
        "og_title": og_title,
        "og_description": og_desc,
        "og_image": og_image,
        "twitter_card": twitter_card,
        "page_type": page_type,
        "is_indexable": is_indexable,
        "depth": depth,
        "status_code": response.status_code,
        "response_time_ms": round(response.elapsed.total_seconds() * 1000),
    }


def crawl_site(start_url, max_pages=MAX_PAGES, max_depth=MAX_DEPTH):
    """BFS crawl of the site, returns list of page data dicts."""
    parsed = urlparse(start_url)
    base_domain = parsed.netloc.lower()
    visited = set()
    queue = deque([(normalize_url(start_url), 0)])
    pages = []

    session = requests.Session()
    session.headers.update(HEADERS)

    print(f"🔍 Starting crawl of {start_url}...")

    while queue and len(pages) < max_pages:
        url, depth = queue.popleft()
        if url in visited or depth > max_depth:
            continue
        visited.add(url)

        try:
            resp = session.get(url, timeout=TIMEOUT, allow_redirects=True)
            ct = resp.headers.get("content-type", "")
            if "text/html" not in ct:
                continue
            page_data = extract_page_data(resp, depth, base_domain)
            pages.append(page_data)
            print(f"  ✅ [{len(pages)}/{max_pages}] {url} ({resp.status_code}, {page_data['response_time_ms']}ms)")

            # Queue internal links
            for link in page_data["internal_links"]:
                if link not in visited:
                    queue.append((link, depth + 1))
        except Exception as e:
            print(f"  ❌ Failed: {url} → {e}")

    print(f"🏁 Crawl complete. {len(pages)} pages analyzed.\n")
    return pages


# ══════════════════════════════════════════════════════════════════════════════
#                          SEO AUDIT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def audit_page(page):
    """Generate page-level audit issues with severity."""
    issues = []
    title = page.get("title", "")
    desc = page.get("description", "")
    h1s = page.get("headings", {}).get("h1", [])
    wc = page.get("word_count", 0)

    # Title checks
    if not title:
        issues.append({"title": "Missing Title Tag", "severity": "High",
                        "why": "Search engines use the title tag as the main clickable headline in search results.",
                        "fix": "Add a descriptive title tag (50-60 characters) containing your target keyword.",
                        "impact": "Directly affects click-through rates and rankings."})
    elif len(title) < 30:
        issues.append({"title": f"Title Too Short ({len(title)} chars)", "severity": "Medium",
                        "why": "Short titles waste valuable SERP real estate and miss keyword opportunities.",
                        "fix": f"Expand the title to 50-60 characters. Current: '{title}'",
                        "impact": "Better titles can increase CTR by 20-30%."})
    elif len(title) > 65:
        issues.append({"title": f"Title Too Long ({len(title)} chars)", "severity": "Low",
                        "why": "Long titles get truncated in search results, cutting off your message.",
                        "fix": "Trim to 50-60 characters while keeping the primary keyword.",
                        "impact": "Prevents message truncation in search results."})

    # Meta description checks
    if not desc or desc == "Not Found":
        issues.append({"title": "Missing Meta Description", "severity": "High",
                        "why": "Without a meta description, Google auto-generates one from random page text.",
                        "fix": "Write a compelling 120-160 character meta description with a call-to-action.",
                        "impact": "Directly improves organic click-through rates."})
    elif len(desc) < 80:
        issues.append({"title": f"Meta Description Too Short ({len(desc)} chars)", "severity": "Medium",
                        "why": "Short descriptions don't give users enough reason to click.",
                        "fix": "Expand to 120-160 characters with benefits and a CTA.",
                        "impact": "Increases search result click-through rates."})
    elif len(desc) > 165:
        issues.append({"title": f"Meta Description Too Long ({len(desc)} chars)", "severity": "Low",
                        "why": "Long descriptions get truncated, potentially cutting off your call-to-action.",
                        "fix": "Trim to 120-160 characters, keeping the main value proposition upfront.",
                        "impact": "Improves messaging clarity in search results."})

    # H1 checks
    if len(h1s) == 0:
        issues.append({"title": "Missing H1 Tag", "severity": "High",
                        "why": "The H1 tag tells search engines the primary topic of the page.",
                        "fix": "Add exactly one descriptive H1 tag containing the main target keyword.",
                        "impact": "Improves content relevance and keyword targeting."})
    elif len(h1s) > 1:
        issues.append({"title": f"Multiple H1 Tags ({len(h1s)} found)", "severity": "Medium",
                        "why": "Multiple H1s dilute the topical focus and confuse search engines.",
                        "fix": "Keep exactly one H1 and change the rest to H2 or H3 tags.",
                        "impact": "Clarifies page topic for better rankings."})

    # Word count
    if wc < 100:
        issues.append({"title": f"Very Low Word Count ({wc} words)", "severity": "High",
                        "why": "Pages with very little content rarely rank because search engines lack context.",
                        "fix": "Add at least 300+ words of valuable, informative content.",
                        "impact": "Boosts topical authority and long-tail keyword visibility."})
    elif wc < 300:
        issues.append({"title": f"Low Word Count ({wc} words)", "severity": "Medium",
                        "why": "Thin content provides insufficient information for ranking well.",
                        "fix": "Expand content to 500+ words with detailed, useful information.",
                        "impact": "Better content depth improves ranking potential."})

    # Canonical
    if not page.get("canonical_url"):
        issues.append({"title": "Missing Canonical Tag", "severity": "Medium",
                        "why": "Without a canonical tag, search engines may treat URL variations as duplicate content.",
                        "fix": "Add a self-referencing canonical tag to every page.",
                        "impact": "Prevents duplicate content issues and link equity dilution."})

    # Viewport
    if not page.get("has_viewport"):
        issues.append({"title": "Missing Viewport Meta Tag", "severity": "High",
                        "why": "Without a viewport tag, the page won't render properly on mobile devices.",
                        "fix": "Add: <meta name='viewport' content='width=device-width, initial-scale=1'>",
                        "impact": "Essential for Google's mobile-first indexing."})

    # Indexing
    if not page.get("is_indexable"):
        issues.append({"title": "Page Blocked from Indexing (noindex)", "severity": "High",
                        "why": "This page has a noindex directive — search engines will ignore it.",
                        "fix": "Remove the noindex directive if this page should appear in search results.",
                        "impact": "Required for the page to receive any organic traffic."})

    # Images
    missing = page.get("missing_alt_images", [])
    total_img = page.get("total_images", 0)
    if total_img > 0 and len(missing) > 0:
        pct = round((len(missing) / total_img) * 100)
        issues.append({"title": f"Images Missing Alt Text ({len(missing)}/{total_img})", "severity": "Medium",
                        "why": "Images without alt text hurt accessibility and miss image search opportunities.",
                        "fix": "Add descriptive alt text to all images describing what they show.",
                        "impact": "Improves accessibility compliance and image search visibility."})

    # Internal links
    il = len(page.get("internal_links", []))
    if il < 3:
        issues.append({"title": f"Very Few Internal Links ({il})", "severity": "Medium",
                        "why": "Pages with few internal links trap PageRank and limit crawl discovery.",
                        "fix": "Add contextual internal links to related services and content pages.",
                        "impact": "Improves crawlability and distributes authority."})

    # Schema
    if not page.get("schema_types"):
        issues.append({"title": "No Structured Data (Schema Markup)", "severity": "Low",
                        "why": "Structured data helps search engines understand your content and can enable rich snippets.",
                        "fix": "Add JSON-LD schema markup (Organization, LocalBusiness, BreadcrumbList, etc.).",
                        "impact": "Can generate rich results, increasing visibility and CTR."})

    # Open Graph
    if not page.get("og_title"):
        issues.append({"title": "Missing Open Graph Tags", "severity": "Low",
                        "why": "Without OG tags, social media platforms generate poor previews of shared links.",
                        "fix": "Add og:title, og:description, og:image, og:url meta tags.",
                        "impact": "Improves social media sharing appearance and engagement."})

    return issues


def calculate_page_score(page):
    """Calculate an overall SEO score (0-100) for a page."""
    score = 100
    title = page.get("title", "")
    desc = page.get("description", "")
    h1s = page.get("headings", {}).get("h1", [])
    wc = page.get("word_count", 0)

    # Title (20 pts)
    if not title:
        score -= 20
    elif len(title) < 30 or len(title) > 70:
        score -= 8

    # Meta description (15 pts)
    if not desc:
        score -= 15
    elif len(desc) < 80 or len(desc) > 170:
        score -= 6

    # H1 (15 pts)
    if len(h1s) == 0:
        score -= 15
    elif len(h1s) > 1:
        score -= 7

    # Word count (10 pts)
    if wc < 100:
        score -= 10
    elif wc < 300:
        score -= 5

    # Canonical (10 pts)
    if not page.get("canonical_url"):
        score -= 10

    # Viewport (10 pts)
    if not page.get("has_viewport"):
        score -= 10

    # Images (5 pts)
    total_img = page.get("total_images", 0)
    missing = len(page.get("missing_alt_images", []))
    if total_img > 0 and missing > 0:
        score -= min(5, round((missing / total_img) * 5))

    # Internal links (5 pts)
    if len(page.get("internal_links", [])) < 3:
        score -= 5

    # Schema (5 pts)
    if not page.get("schema_types"):
        score -= 3

    # OG tags (5 pts)
    if not page.get("og_title"):
        score -= 2

    return max(0, min(100, score))


def generate_sitewide_summary(pages):
    """Aggregate stats across all pages."""
    total = len(pages)
    if total == 0:
        return {}

    titles_present = sum(1 for p in pages if p.get("title"))
    descs_present = sum(1 for p in pages if p.get("description"))
    h1_correct = sum(1 for p in pages if len(p.get("headings", {}).get("h1", [])) == 1)
    viewports = sum(1 for p in pages if p.get("has_viewport"))
    canonicals = sum(1 for p in pages if p.get("canonical_url"))
    indexable = sum(1 for p in pages if p.get("is_indexable"))
    has_schema = sum(1 for p in pages if p.get("schema_types"))
    has_og = sum(1 for p in pages if p.get("og_title"))

    all_titles = [p.get("title", "") for p in pages if p.get("title")]
    unique_titles = len(set(all_titles))
    all_descs = [p.get("description", "") for p in pages if p.get("description")]
    unique_descs = len(set(all_descs))

    avg_wc = round(sum(p.get("word_count", 0) for p in pages) / total)
    avg_response = round(sum(p.get("response_time_ms", 0) for p in pages) / total)

    total_images = sum(p.get("total_images", 0) for p in pages)
    total_missing_alt = sum(len(p.get("missing_alt_images", [])) for p in pages)

    scores = [calculate_page_score(p) for p in pages]
    avg_score = round(sum(scores) / total)

    return {
        "total_pages": total,
        "avg_score": avg_score,
        "titles_present": titles_present,
        "titles_pct": round(titles_present / total * 100),
        "descs_present": descs_present,
        "descs_pct": round(descs_present / total * 100),
        "unique_titles": unique_titles,
        "unique_descs": unique_descs,
        "h1_correct": h1_correct,
        "h1_pct": round(h1_correct / total * 100),
        "viewport_pct": round(viewports / total * 100),
        "canonical_pct": round(canonicals / total * 100),
        "indexable_pct": round(indexable / total * 100),
        "schema_pct": round(has_schema / total * 100),
        "og_pct": round(has_og / total * 100),
        "avg_word_count": avg_wc,
        "avg_response_ms": avg_response,
        "total_images": total_images,
        "missing_alt_images": total_missing_alt,
        "alt_coverage_pct": round(((total_images - total_missing_alt) / max(total_images, 1)) * 100),
        "page_scores": scores,
    }


# ══════════════════════════════════════════════════════════════════════════════
#                     WORD DOCUMENT GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)  # slate-900
    return h


def severity_color(sev):
    if sev == "High":
        return RGBColor(0xDC, 0x26, 0x26)  # red
    elif sev == "Medium":
        return RGBColor(0xF5, 0x9E, 0x0B)  # amber
    return RGBColor(0x64, 0x74, 0x8B)  # slate


def score_label(s):
    if s >= 80:
        return "Excellent", "22C55E"
    elif s >= 60:
        return "Good", "84CC16"
    elif s >= 40:
        return "Needs Attention", "F59E0B"
    return "Poor", "EF4444"


def add_metric_row(table, label, value, benchmark, status):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)
    row.cells[2].text = benchmark
    row.cells[3].text = status

    # Color the status cell
    if "✅" in status or "Pass" in status:
        set_cell_shading(row.cells[3], "DCFCE7")
    elif "⚠️" in status or "Warning" in status:
        set_cell_shading(row.cells[3], "FEF9C3")
    elif "❌" in status or "Fail" in status:
        set_cell_shading(row.cells[3], "FEE2E2")


def generate_docx_report(pages, target_url, output_path):
    """Generate a professional Word document SEO audit report."""
    doc = Document()

    # ── Page Setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Heading styles
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    summary = generate_sitewide_summary(pages)
    overall_score = summary.get("avg_score", 0)
    label, color_hex = score_label(overall_score)

    # ══════════════════════════════════════════════════════════════════════════
    #                        COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("SEO AUDIT REPORT")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_sub.add_run("Comprehensive Technical & On-Page Analysis")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("")

    cover_url = doc.add_paragraph()
    cover_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_url.add_run(target_url)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x49, 0x46, 0xF5)

    doc.add_paragraph("")
    doc.add_paragraph("")

    cover_date = doc.add_paragraph()
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_date.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    cover_score = doc.add_paragraph()
    cover_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_score.add_run(f"Overall SEO Score: {overall_score}/100 — {label}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Sitewide Health Overview",
        "3. Technical SEO Metrics",
        "4. Content Analysis",
        "5. Page-by-Page Audit",
        "6. Issue Summary & Priority Matrix",
        "7. Recommendations & Action Plan",
        "8. Appendix: Raw Data",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "1. Executive Summary", level=1)

    p = doc.add_paragraph()
    p.add_run(f"This report provides a comprehensive SEO audit of ").font.size = Pt(10)
    run = p.add_run(target_url)
    run.bold = True
    run.font.color.rgb = RGBColor(0x49, 0x46, 0xF5)
    p.add_run(f", covering {summary['total_pages']} pages analyzed on {datetime.now().strftime('%B %d, %Y')}.")

    doc.add_paragraph("")

    # Key findings box
    key_findings = doc.add_paragraph()
    key_findings.add_run("Key Findings:").bold = True
    
    high_issues = sum(1 for p in pages for iss in audit_page(p) if iss["severity"] == "High")
    med_issues = sum(1 for p in pages for iss in audit_page(p) if iss["severity"] == "Medium")
    low_issues = sum(1 for p in pages for iss in audit_page(p) if iss["severity"] == "Low")

    findings = [
        f"Overall SEO Score: {overall_score}/100 ({label})",
        f"Pages Analyzed: {summary['total_pages']}",
        f"Average Page Load Time: {summary['avg_response_ms']}ms",
        f"High Priority Issues: {high_issues}",
        f"Medium Priority Issues: {med_issues}",
        f"Low Priority Issues: {low_issues}",
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   2. SITEWIDE HEALTH OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "2. Sitewide Health Overview", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Current Value"
    hdr[2].text = "Benchmark"
    hdr[3].text = "Status"
    for cell in hdr:
        set_cell_shading(cell, "1E293B")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)

    tp = summary["titles_pct"]
    add_metric_row(table, "Title Tag Coverage",
                   f"{summary['titles_present']}/{summary['total_pages']} ({tp}%)",
                   "100%", "✅ Pass" if tp == 100 else "❌ Fail")

    dp = summary["descs_pct"]
    add_metric_row(table, "Meta Description Coverage",
                   f"{summary['descs_present']}/{summary['total_pages']} ({dp}%)",
                   "100%", "✅ Pass" if dp == 100 else "❌ Fail")

    hp = summary["h1_pct"]
    add_metric_row(table, "Correct H1 Usage (1 per page)",
                   f"{summary['h1_correct']}/{summary['total_pages']} ({hp}%)",
                   "100%", "✅ Pass" if hp == 100 else ("⚠️ Warning" if hp >= 50 else "❌ Fail"))

    add_metric_row(table, "Mobile Viewport Tag",
                   f"{summary['viewport_pct']}%",
                   "100%", "✅ Pass" if summary["viewport_pct"] == 100 else "❌ Fail")

    add_metric_row(table, "Canonical Tag Coverage",
                   f"{summary['canonical_pct']}%",
                   "100%", "✅ Pass" if summary["canonical_pct"] == 100 else "⚠️ Warning")

    add_metric_row(table, "Indexable Pages",
                   f"{summary['indexable_pct']}%",
                   "100%", "✅ Pass" if summary["indexable_pct"] == 100 else "⚠️ Warning")

    add_metric_row(table, "Image Alt Text Coverage",
                   f"{summary['alt_coverage_pct']}%",
                   "100%", "✅ Pass" if summary["alt_coverage_pct"] >= 90 else ("⚠️ Warning" if summary["alt_coverage_pct"] >= 50 else "❌ Fail"))

    add_metric_row(table, "Structured Data (Schema)",
                   f"{summary['schema_pct']}%",
                   "Recommended", "✅ Pass" if summary["schema_pct"] >= 50 else "⚠️ Warning")

    add_metric_row(table, "Open Graph Tags",
                   f"{summary['og_pct']}%",
                   "Recommended", "✅ Pass" if summary["og_pct"] >= 50 else "⚠️ Warning")

    add_metric_row(table, "Unique Title Tags",
                   f"{summary['unique_titles']}/{summary['titles_present']}",
                   "All unique", "✅ Pass" if summary["unique_titles"] == summary["titles_present"] else "⚠️ Warning")

    add_metric_row(table, "Unique Meta Descriptions",
                   f"{summary['unique_descs']}/{summary['descs_present']}",
                   "All unique", "✅ Pass" if summary["unique_descs"] == summary["descs_present"] else "⚠️ Warning")

    doc.add_paragraph("")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   3. TECHNICAL SEO METRICS
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "3. Technical SEO Metrics", level=1)

    add_styled_heading(doc, "3.1 Page Response Times", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Average Response Time: ").bold = True
    p.add_run(f"{summary['avg_response_ms']}ms")
    if summary['avg_response_ms'] < 500:
        p.add_run(" — ✅ Excellent").font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
    elif summary['avg_response_ms'] < 1500:
        p.add_run(" — ⚠️ Acceptable").font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
    else:
        p.add_run(" — ❌ Slow").font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

    # Response time table
    rt_table = doc.add_table(rows=1, cols=3)
    rt_table.style = 'Light Grid Accent 1'
    rt_hdr = rt_table.rows[0].cells
    rt_hdr[0].text = "Page URL"
    rt_hdr[1].text = "Status Code"
    rt_hdr[2].text = "Response Time"
    for cell in rt_hdr:
        set_cell_shading(cell, "1E293B")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)

    for page in sorted(pages, key=lambda x: x.get("response_time_ms", 0), reverse=True):
        row = rt_table.add_row()
        url_short = page["url"].replace(target_url, "/") or "/"
        row.cells[0].text = url_short[:60]
        row.cells[1].text = str(page.get("status_code", "N/A"))
        row.cells[2].text = f"{page.get('response_time_ms', 'N/A')}ms"
        for cell_para in row.cells[0].paragraphs:
            for run in cell_para.runs:
                run.font.size = Pt(8)

    doc.add_paragraph("")

    add_styled_heading(doc, "3.2 Structured Data Found", level=2)
    all_schemas = set()
    for page in pages:
        all_schemas.update(page.get("schema_types", []))
    if all_schemas:
        for s in sorted(all_schemas):
            doc.add_paragraph(f"• {s}", style="List Bullet")
    else:
        doc.add_paragraph("❌ No structured data (JSON-LD, Microdata) was detected on any page.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   4. CONTENT ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "4. Content Analysis", level=1)

    add_styled_heading(doc, "4.1 Word Count Distribution", level=2)
    wc_table = doc.add_table(rows=1, cols=4)
    wc_table.style = 'Light Grid Accent 1'
    wc_hdr = wc_table.rows[0].cells
    wc_hdr[0].text = "Page"
    wc_hdr[1].text = "Word Count"
    wc_hdr[2].text = "Page Type"
    wc_hdr[3].text = "Assessment"
    for cell in wc_hdr:
        set_cell_shading(cell, "1E293B")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)

    for page in pages:
        wc = page.get("word_count", 0)
        row = wc_table.add_row()
        url_short = page["url"].replace(target_url, "/") or "/"
        row.cells[0].text = url_short[:55]
        row.cells[1].text = str(wc)
        row.cells[2].text = page.get("page_type", "Other")
        if wc >= 300:
            row.cells[3].text = "✅ Good"
            set_cell_shading(row.cells[3], "DCFCE7")
        elif wc >= 100:
            row.cells[3].text = "⚠️ Thin"
            set_cell_shading(row.cells[3], "FEF9C3")
        else:
            row.cells[3].text = "❌ Very Thin"
            set_cell_shading(row.cells[3], "FEE2E2")
        for cell_para in row.cells[0].paragraphs:
            for run in cell_para.runs:
                run.font.size = Pt(8)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(f"Average Word Count: ").bold = True
    p.add_run(f"{summary['avg_word_count']} words")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   5. PAGE-BY-PAGE AUDIT
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "5. Page-by-Page Audit", level=1)

    for idx, page in enumerate(pages, 1):
        score = calculate_page_score(page)
        label_txt, clr = score_label(score)
        issues = audit_page(page)

        add_styled_heading(doc, f"5.{idx} {page['url']}", level=2)

        # Score badge
        p = doc.add_paragraph()
        run = p.add_run(f"SEO Score: {score}/100 — {label_txt}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(int(clr[:2], 16), int(clr[2:4], 16), int(clr[4:], 16))

        # Page info table
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ("Page Type", page.get("page_type", "N/A")),
            ("Title", page.get("title", "— Missing —") or "— Missing —"),
            ("Title Length", f"{len(page.get('title', ''))} characters"),
            ("Meta Description", (page.get("description", "")[:100] + "...") if len(page.get("description", "")) > 100 else (page.get("description", "") or "— Missing —")),
            ("Meta Description Length", f"{len(page.get('description', ''))} characters"),
            ("Word Count", str(page.get("word_count", 0))),
            ("H1 Tags", ", ".join(page.get("headings", {}).get("h1", [])) or "— None —"),
            ("H1 Count", str(len(page.get("headings", {}).get("h1", [])))),
            ("Canonical URL", page.get("canonical_url", "") or "— Missing —"),
            ("Viewport Tag", "✅ Present" if page.get("has_viewport") else "❌ Missing"),
            ("Indexable", "✅ Yes" if page.get("is_indexable") else "❌ No (noindex)"),
            ("Images", f"{page.get('total_images', 0)} total, {len(page.get('missing_alt_images', []))} missing alt"),
            ("Internal Links", str(len(page.get("internal_links", [])))),
            ("External Links", str(len(page.get("external_links", [])))),
            ("Dofollow / Nofollow", f"{page.get('dofollow_links', 0)} / {page.get('nofollow_links', 0)}"),
            ("Schema Types", ", ".join(page.get("schema_types", [])) or "None"),
            ("Response Time", f"{page.get('response_time_ms', 'N/A')}ms"),
        ]

        for label_name, val in info_data:
            row = info_table.add_row()
            row.cells[0].text = label_name
            row.cells[1].text = val
            set_cell_shading(row.cells[0], "F1F5F9")
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

        doc.add_paragraph("")

        # Headings structure
        add_styled_heading(doc, "Heading Structure", level=3)
        headings = page.get("headings", {})
        has_headings = False
        for level_num in range(1, 7):
            tag = f"h{level_num}"
            tag_list = headings.get(tag, [])
            if tag_list:
                has_headings = True
                p = doc.add_paragraph()
                run = p.add_run(f"H{level_num} ({len(tag_list)}): ")
                run.bold = True
                run.font.size = Pt(9)
                for h_text in tag_list[:5]:
                    doc.add_paragraph(f'  "{h_text}"', style='List Bullet')
                if len(tag_list) > 5:
                    doc.add_paragraph(f"  ... and {len(tag_list) - 5} more")
        if not has_headings:
            doc.add_paragraph("❌ No heading tags found on this page.")

        # Issues for this page
        if issues:
            add_styled_heading(doc, "Issues Found", level=3)
            for iss in issues:
                p = doc.add_paragraph()
                sev_run = p.add_run(f"[{iss['severity']}] ")
                sev_run.bold = True
                sev_run.font.color.rgb = severity_color(iss["severity"])
                title_run = p.add_run(iss["title"])
                title_run.bold = True

                why_p = doc.add_paragraph()
                why_p.paragraph_format.left_indent = Cm(1)
                why_run = why_p.add_run("Why It Matters: ")
                why_run.bold = True
                why_run.font.size = Pt(9)
                why_p.add_run(iss["why"]).font.size = Pt(9)

                fix_p = doc.add_paragraph()
                fix_p.paragraph_format.left_indent = Cm(1)
                fix_run = fix_p.add_run("Recommended Fix: ")
                fix_run.bold = True
                fix_run.font.size = Pt(9)
                fix_p.add_run(iss["fix"]).font.size = Pt(9)

                imp_p = doc.add_paragraph()
                imp_p.paragraph_format.left_indent = Cm(1)
                imp_run = imp_p.add_run("SEO Impact: ")
                imp_run.bold = True
                imp_run.font.size = Pt(9)
                imp_run.font.color.rgb = RGBColor(0x49, 0x46, 0xF5)
                imp_p.add_run(iss["impact"]).font.size = Pt(9)

                doc.add_paragraph("")  # spacer
        else:
            p = doc.add_paragraph()
            run = p.add_run("✅ No issues found — this page is well optimized!")
            run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
            run.bold = True

        if idx < len(pages):
            doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                 6. ISSUE SUMMARY & PRIORITY MATRIX
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "6. Issue Summary & Priority Matrix", level=1)

    # Collect all issues across pages
    all_issues = {}
    for page in pages:
        for iss in audit_page(page):
            key = iss["title"]
            if key not in all_issues:
                all_issues[key] = {"severity": iss["severity"], "count": 0, "pages": [],
                                   "why": iss["why"], "fix": iss["fix"], "impact": iss["impact"]}
            all_issues[key]["count"] += 1
            url_short = page["url"].replace(target_url, "/") or "/"
            all_issues[key]["pages"].append(url_short[:40])

    if all_issues:
        # Sort: High first, then Medium, then Low
        sev_order = {"High": 0, "Medium": 1, "Low": 2}
        sorted_issues = sorted(all_issues.items(), key=lambda x: (sev_order.get(x[1]["severity"], 3), -x[1]["count"]))

        issue_table = doc.add_table(rows=1, cols=4)
        issue_table.style = 'Light Grid Accent 1'
        ihdr = issue_table.rows[0].cells
        ihdr[0].text = "Issue"
        ihdr[1].text = "Severity"
        ihdr[2].text = "Pages Affected"
        ihdr[3].text = "Recommended Fix"
        for cell in ihdr:
            set_cell_shading(cell, "1E293B")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
                    run.font.size = Pt(9)

        for issue_name, data in sorted_issues:
            row = issue_table.add_row()
            row.cells[0].text = issue_name
            row.cells[1].text = data["severity"]
            row.cells[2].text = f"{data['count']} page(s)"
            row.cells[3].text = data["fix"][:120]

            # Color severity cell
            if data["severity"] == "High":
                set_cell_shading(row.cells[1], "FEE2E2")
            elif data["severity"] == "Medium":
                set_cell_shading(row.cells[1], "FEF9C3")
            else:
                set_cell_shading(row.cells[1], "E0E7FF")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    else:
        p = doc.add_paragraph()
        run = p.add_run("✅ No issues found across any pages!")
        run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
        run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   7. RECOMMENDATIONS & ACTION PLAN
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "7. Recommendations & Action Plan", level=1)

    p = doc.add_paragraph()
    p.add_run("Based on the audit findings, here is a prioritized action plan:").font.size = Pt(10)
    doc.add_paragraph("")

    # High Priority
    high_issues_list = [(k, v) for k, v in sorted_issues if v["severity"] == "High"] if all_issues else []
    if high_issues_list:
        add_styled_heading(doc, "🔴 High Priority (Fix Immediately)", level=2)
        for i, (name, data) in enumerate(high_issues_list, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {name}")
            run.bold = True
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.left_indent = Cm(1)
            desc_p.add_run(f"Affects {data['count']} page(s). ").font.size = Pt(9)
            desc_p.add_run(data["fix"]).font.size = Pt(9)
            doc.add_paragraph("")

    # Medium Priority
    med_issues_list = [(k, v) for k, v in sorted_issues if v["severity"] == "Medium"] if all_issues else []
    if med_issues_list:
        add_styled_heading(doc, "🟡 Medium Priority (Fix Soon)", level=2)
        for i, (name, data) in enumerate(med_issues_list, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {name}")
            run.bold = True
            run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)

            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.left_indent = Cm(1)
            desc_p.add_run(f"Affects {data['count']} page(s). ").font.size = Pt(9)
            desc_p.add_run(data["fix"]).font.size = Pt(9)
            doc.add_paragraph("")

    # Low Priority
    low_issues_list = [(k, v) for k, v in sorted_issues if v["severity"] == "Low"] if all_issues else []
    if low_issues_list:
        add_styled_heading(doc, "🟢 Low Priority (Nice to Have)", level=2)
        for i, (name, data) in enumerate(low_issues_list, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {name}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.left_indent = Cm(1)
            desc_p.add_run(f"Affects {data['count']} page(s). ").font.size = Pt(9)
            desc_p.add_run(data["fix"]).font.size = Pt(9)
            doc.add_paragraph("")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #                   8. APPENDIX: RAW DATA
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "8. Appendix: Page Score Summary", level=1)

    score_table = doc.add_table(rows=1, cols=4)
    score_table.style = 'Light Grid Accent 1'
    shdr = score_table.rows[0].cells
    shdr[0].text = "Page URL"
    shdr[1].text = "Score"
    shdr[2].text = "Status"
    shdr[3].text = "Issues Count"
    for cell in shdr:
        set_cell_shading(cell, "1E293B")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)

    for page in pages:
        score = calculate_page_score(page)
        issues = audit_page(page)
        label_txt, clr = score_label(score)
        row = score_table.add_row()
        url_short = page["url"].replace(target_url, "/") or "/"
        row.cells[0].text = url_short[:55]
        row.cells[1].text = f"{score}/100"
        row.cells[2].text = label_txt
        row.cells[3].text = str(len(issues))

        if score >= 80:
            set_cell_shading(row.cells[1], "DCFCE7")
        elif score >= 60:
            set_cell_shading(row.cells[1], "FEF9C3")
        else:
            set_cell_shading(row.cells[1], "FEE2E2")

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph("")

    # External domains found
    add_styled_heading(doc, "External Domains Linked", level=2)
    all_ext = set()
    for page in pages:
        all_ext.update(page.get("external_domains", []))
    if all_ext:
        for d in sorted(all_ext)[:30]:
            doc.add_paragraph(f"• {d}", style="List Bullet")
        if len(all_ext) > 30:
            doc.add_paragraph(f"... and {len(all_ext) - 30} more domains")
    else:
        doc.add_paragraph("No external domains linked.")

    doc.add_paragraph("")

    # Footer
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("— End of Report —")
    run.bold = True
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.size = Pt(10)

    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = footer_p2.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    run2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run2.font.size = Pt(8)

    # Save
    doc.save(output_path)
    print(f"\n📄 Report saved to: {output_path}")
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
#                              MAIN
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 70)
    print("   SEO AUDIT REPORT GENERATOR")
    print(f"   Target: {TARGET_URL}")
    print("=" * 70)
    print()

    start = time.time()
    pages = crawl_site(TARGET_URL)

    if not pages:
        print("❌ No pages could be crawled. Please check the URL.")
        exit(1)

    # Generate report
    output_file = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"SEO_Audit_Report_{urlparse(TARGET_URL).netloc}_{datetime.now().strftime('%Y%m%d')}.docx"
    )
    generate_docx_report(pages, TARGET_URL, output_file)

    elapsed = round(time.time() - start, 1)
    print(f"\n✅ Complete! Generated in {elapsed} seconds.")
    print(f"📂 File: {output_file}")
